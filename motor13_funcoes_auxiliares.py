import fitz
import re
import pandas as pd
from docx import Document

def extrair_dados_processo(doc):
    texto = doc[0].get_text()
    dados = {}

    # Buscar número do processo (formato CNJ)
    match_processo = re.search(r'\d{7}-\d{2}\.\d{4}\.\d{1,2}\.\d{4}', texto)
    dados["Número do Processo"] = match_processo.group(0) if match_processo else "Não localizado"

    # Buscar Reclamante (várias possibilidades)
    match_reclamante = re.search(r'(Reclamante|RECLAMANTE|RECLAMADO):?\s*(.+)', texto)
    if match_reclamante:
        dados["Reclamante"] = match_reclamante.group(2).split("\n")[0].strip()
    else:
        dados["Reclamante"] = "Não localizado"

    # Buscar Reclamada
    match_reclamada = re.search(r'(Reclamada|RECLAMADA|RECLAMADO):?\s*(.+)', texto)
    if match_reclamada:
        dados["Reclamada"] = match_reclamada.group(2).split("\n")[0].strip()
    else:
        dados["Reclamada"] = "Não localizado"

    # Buscar Data de Autuação
    match_data = re.search(r'(\d{2}/\d{2}/\d{4})', texto)
    dados["Data de Autuação"] = match_data.group(0) if match_data else "Não localizada"

    # Buscar Valor da Causa
    match_valor = re.search(r'Valor da Causa[:\s]*R?\$?\s*([\d\.,]+)', texto)
    dados["Valor da Causa"] = match_valor.group(1) if match_valor else "Não localizado"

    return dados

def extrair_sumario(doc):
    total_paginas = len(doc)
    for i in range(total_paginas - 1, max(0, total_paginas - 10), -1):
        if "SUMÁRIO" in doc[i].get_text():
            return i, doc[i].get_text()
    return -1, ""

def reconstruir_sumario_completo(doc, pagina_inicial_sumario):
    linhas = []
    for i in range(pagina_inicial_sumario, len(doc)):
        linhas.extend(doc[i].get_text().splitlines())

    blocos = []
    i = 0
    while i < len(linhas):
        if re.fullmatch(r"[a-fA-F0-9]{7}", linhas[i].strip()):
            blocos.append(linhas[i:i+4])
            i += 4
        else:
            i += 1

    df = pd.DataFrame(blocos, columns=["ID", "Data", "Documento", "Tipo"])
    df["ID Posterior"] = df["ID"].shift(-1)
    return df

def criar_indexador(df, doc, pagina_sumario):
    paginas_por_id = {}
    for i in range(pagina_sumario):
        texto = doc[i].get_text()
        for idx in df["ID"]:
            if idx in texto:
                if idx not in paginas_por_id:
                    paginas_por_id[idx] = []
                paginas_por_id[idx].append(i + 1)

    paginas_indexadas = []
    for idx in df["ID"]:
        pags = paginas_por_id.get(idx, [])
        if pags:
            paginas_indexadas.append((min(pags), max(pags)))
        else:
            paginas_indexadas.append((None, None))

    df["Página Inicial"], df["Página Final"] = zip(*paginas_indexadas)
    return df

def localizar_trct(df, doc, reclamante, pdf_path):
    for _, row in df.iterrows():
        if any(x in row["Tipo"].lower() for x in ["trct", "rescisão", "rescisao"]):
            inicio, fim = int(row["Página Inicial"]) - 1, int(row["Página Final"]) - 1
            conteudo = "\n".join([doc[i].get_text() for i in range(inicio, fim + 1)])
            return {
                "Texto": conteudo,
                "ID": row["ID"],
                "Página Inicial": row["Página Inicial"],
                "Página Final": row["Página Final"]
            }
    return {"Texto": "", "ID": "", "Página Inicial": "", "Página Final": ""}

def gerar_dados_para_calculo(resposta):
    # Simulação simples
    return pd.DataFrame([{"Verba": "Horas Extras", "Valor": "R$ 2.000,00"}])

def gerar_relatorio(dados_processo, dados_trct, resposta, sumario_filtrado, df_resultado):
    doc = Document()
    doc.add_heading("RELATÓRIO PERICIAL - MOTOR 13", level=1)
    doc.add_heading("Dados do Processo", level=2)
    for chave, valor in dados_processo.items():
        doc.add_paragraph(f"{chave}: {valor}")

    doc.add_heading("TRCT (Resumo)", level=2)
    if dados_trct["Texto"]:
        doc.add_paragraph(f"Páginas: {dados_trct['Página Inicial']} a {dados_trct['Página Final']}")
    else:
        doc.add_paragraph("TRCT não localizado.")

    doc.add_heading("Sumário (Filtrado)", level=2)
    for _, row in sumario_filtrado.iterrows():
        doc.add_paragraph(f"{row['Tipo']} - {row['Documento']} (ID: {row['ID']})")

    doc.add_heading("Análise das Decisões", level=2)
    doc.add_paragraph(resposta)

    doc.add_heading("Resumo das Verbas", level=2)
    for _, row in df_resultado.iterrows():
        doc.add_paragraph(f"{row['Verba']}: {row['Valor']}")

    return doc